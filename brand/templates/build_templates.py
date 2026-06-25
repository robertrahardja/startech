"""Generate Startech .docx templates from the brand guidelines.

Produces:
  - Startech-Quotation-Template.docx
  - Startech-Invoice-Template.docx

Run: python3 build_templates.py
Requires: python-docx >= 1.0
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

HERE = Path(__file__).parent
LOGO_LIGHT = HERE.parent / "logos" / "startech-full-light.png"
MARK_LIGHT = HERE.parent / "logos" / "startech-mark-light.png"

# ─── Brand tokens (print palette) ─────────────────────────────────────────
BLUE = RGBColor(0x4F, 0x6B, 0xF7)
BLUE_DARK = RGBColor(0x2A, 0x3A, 0x8C)
GOLD = RGBColor(0xD4, 0xA8, 0x53)
TEXT = RGBColor(0x18, 0x18, 0x1B)
MUTED = RGBColor(0x71, 0x71, 0x7A)
HAIRLINE = "E4E4E7"
BLUE_HEX = "4F6BF7"
ALT_ROW_HEX = "F5F5F8"

FONT_SANS = "Inter"
FONT_SANS_FALLBACK = "Calibri"
FONT_SERIF = "Georgia"


# ─── Low-level XML helpers ────────────────────────────────────────────────
def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side, spec in (
        ("top", top), ("left", left), ("bottom", bottom), ("right", right)
    ):
        if spec is None:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))  # 1/8pt
        el.set(qn("w:color"), spec.get("color", "auto"))
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def remove_all_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tbl_borders.append(el)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_borders)


def add_horizontal_rule(paragraph, *, color_hex: str, size_pt: float = 3):
    """Insert a bottom border on a paragraph to create a horizontal rule."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))  # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ─── Style helpers ────────────────────────────────────────────────────────
def configure_base_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_SANS)
    r_fonts.set(qn("w:hAnsi"), FONT_SANS)
    r_fonts.set(qn("w:cs"), FONT_SANS)
    r_fonts.set(qn("w:eastAsia"), FONT_SANS)

    # Paragraph defaults
    p_fmt = normal.paragraph_format
    p_fmt.space_after = Pt(4)
    p_fmt.line_spacing = 1.35


def run(paragraph, text, *, bold=False, italic=False, size=10,
        color=TEXT, font=FONT_SANS, spacing=None, upper=False):
    r = paragraph.add_run(text.upper() if upper else text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if spacing is not None:
        rpr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing * 20)))  # 1/20 pt
        rpr.append(sp)
    return r


def h2(doc_or_cell, text):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, bold=True, size=11, color=BLUE_DARK,
        spacing=1.2, upper=True)
    return p


def body(container, text, *, bold=False, italic=False, color=TEXT, size=10):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def bullet(container, text):
    p = container.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=10)
    return p


# ─── Page setup ───────────────────────────────────────────────────────────
def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)


# ─── Shared document chrome ───────────────────────────────────────────────
def build_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header

    # Clear default paragraph
    header_para = header.paragraphs[0]
    header_para.text = ""

    # Table: company name | mark logo
    table = header.add_table(rows=1, cols=2, width=Mm(166))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    remove_all_table_borders(table)
    table.columns[0].width = Mm(130)
    table.columns[1].width = Mm(36)

    name_cell = table.cell(0, 0)
    name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    name_para = name_cell.paragraphs[0]
    name_para.paragraph_format.space_after = Pt(2)
    run(name_para, "STARTECH INNOVATION PTE LTD",
        bold=True, size=11, color=BLUE_DARK, spacing=1.6)
    add_horizontal_rule(name_para, color_hex=BLUE_HEX, size_pt=2.5)

    logo_cell = table.cell(0, 1)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if MARK_LIGHT.exists():
        logo_para.add_run().add_picture(str(MARK_LIGHT), height=Mm(11))

    # Spacer paragraph after the table
    spacer = header.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def build_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(
        footer_para,
        "This document contains commercially confidential information and "
        "is intended solely for the named recipient. "
        "Unauthorized distribution is strictly prohibited.",
        italic=True, size=8, color=MUTED,
    )
    # Add hairline above
    add_horizontal_rule(footer_para, color_hex=HAIRLINE, size_pt=0.75)


# ─── Table primitives ─────────────────────────────────────────────────────
def styled_table(doc: Document, headers, rows, col_widths_mm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    remove_all_table_borders(table)

    for i, w in enumerate(col_widths_mm):
        table.columns[i].width = Mm(w)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Mm(col_widths_mm[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, BLUE_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run(p, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF),
            spacing=0.8, upper=True)

    # Body rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.width = Mm(col_widths_mm[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_HEX)
            set_cell_borders(
                cell,
                bottom={"val": "single", "sz": 4, "color": HAIRLINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run(p, val, size=10)

    return table


# ─── Quotation template ───────────────────────────────────────────────────
def build_quotation():
    doc = Document()
    configure_base_styles(doc)
    setup_page(doc)
    build_header(doc)
    build_footer(doc)

    # Title block
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    run(title, "QUOTATION", size=26, color=TEXT, spacing=1.5)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run(subtitle, "[Project / engagement title]", size=11, color=MUTED,
        italic=True)

    # Meta + client block (two columns)
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.autofit = False
    remove_all_table_borders(meta_table)
    meta_table.columns[0].width = Mm(60)
    meta_table.columns[1].width = Mm(106)

    left_cell = meta_table.cell(0, 0)
    left_cell.width = Mm(60)
    body(left_cell, "REFERENCE", bold=True, size=9, color=MUTED)
    body(left_cell, "STI-Q[YYMM]-[DDDD]")
    body(left_cell, "DATE", bold=True, size=9, color=MUTED)
    body(left_cell, "[8 Jan 2026]")
    body(left_cell, "UEN", bold=True, size=9, color=MUTED)
    body(left_cell, "202110461R")

    right_cell = meta_table.cell(0, 1)
    right_cell.width = Mm(106)
    body(right_cell, "PREPARED FOR", bold=True, size=9, color=MUTED)
    body(right_cell, "[CLIENT COMPANY PTE LTD]", bold=True)
    body(right_cell, "[Street address, #floor-unit, Singapore 000000]")
    body(right_cell, "")
    body(right_cell, "ATTENTION", bold=True, size=9, color=MUTED)
    body(right_cell, "[Dr. Recipient Name] — [Title]")

    # Project snapshot
    h2(doc, "Project snapshot")
    snapshot = doc.add_table(rows=3, cols=2)
    snapshot.autofit = False
    remove_all_table_borders(snapshot)
    snapshot.columns[0].width = Mm(40)
    snapshot.columns[1].width = Mm(126)
    for i, (label, value) in enumerate([
        ("Project", "[Short project name]"),
        ("Timeline", "[e.g. End of February 2026 (6 weeks)]"),
        ("Deliverable", "[e.g. Functional demo — Web & Mobile]"),
    ]):
        lc = snapshot.cell(i, 0)
        rc = snapshot.cell(i, 1)
        lc.width = Mm(40)
        rc.width = Mm(126)
        p = lc.paragraphs[0]
        run(p, label, bold=True, size=9, color=MUTED, upper=True, spacing=0.8)
        p2 = rc.paragraphs[0]
        run(p2, value, size=10)

    h2(doc, "Project overview")
    body(
        doc,
        "[Two-to-four-sentence narrative of the engagement. State the "
        "problem, the proposed approach, and the user/beneficiary. Keep "
        "it plain-spoken — no buzzwords.]",
    )

    h2(doc, "Scope of work")

    for part_label in [
        "Part 1 — [Work package name]",
        "Part 2 — [Work package name]",
        "Part 3 — [Work package name]",
    ]:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(8)
        sub.paragraph_format.space_after = Pt(3)
        run(sub, part_label, bold=True, size=11, color=TEXT)
        bullet(doc, "[Deliverable or capability]")
        bullet(doc, "[Deliverable or capability]")
        bullet(doc, "[Deliverable or capability]")

    h2(doc, "Technical specifications")
    body(doc, "Stack: [e.g. React, Hono.js, Cloudflare Workers, D1]")
    body(doc, "Integrations: [e.g. Coursera LMS, OpenAI GPT-4]")

    h2(doc, "Project timeline")
    styled_table(
        doc,
        headers=["Phase", "Duration", "Milestones"],
        rows=[
            ["Week 1–2", "Requirements & design", "[Approval of specs]"],
            ["Week 3–4", "Core development", "[Backend, API, flows]"],
            ["Week 5–6", "Frontend & integration", "[UI, tests, handover]"],
        ],
        col_widths_mm=[30, 46, 90],
    )

    h2(doc, "Investment structure")
    styled_table(
        doc,
        headers=["Component", "Cost"],
        rows=[
            ["Part 1 — [Work package]", "SGD 0,000"],
            ["Part 2 — [Work package]", "SGD 0,000"],
            ["Part 3 — [Work package]", "SGD 0,000"],
            ["TOTAL PROJECT INVESTMENT", "SGD 0,000"],
        ],
        col_widths_mm=[126, 40],
    )

    h2(doc, "Payment schedule")
    bullet(doc, "Initial Payment: SGD [0,000] upon contract signing")
    bullet(doc, "Mid-Project Payment: SGD [0,000] upon [milestone]")
    bullet(doc, "Final Payment: SGD [0,000] upon handover & documentation")
    body(doc, "")
    body(
        doc,
        "Payment terms: invoices payable within 30 days of issue. "
        "Late payment charges: 1.5% per month.",
        italic=True, color=MUTED, size=9,
    )

    h2(doc, "Contact & next steps")
    body(
        doc,
        "Robert Rahardja  |  Managing Director, Startech Innovation Pte Ltd",
        bold=True,
    )
    body(
        doc,
        "Email: robert@startech-innovation.com  |  Direct: +65 9069 3236  "
        "|  UEN: 202110461R",
    )
    body(doc, "Address: 1003 Bukit Merah Central #06-07, Singapore 159836")

    h2(doc, "Agreement")
    body(
        doc,
        "By signing below, both parties acknowledge understanding and "
        "acceptance of all terms, conditions, and specifications outlined "
        "in this quotation.",
    )

    # Signature block (two columns)
    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = False
    remove_all_table_borders(sig)
    sig.columns[0].width = Mm(80)
    sig.columns[1].width = Mm(86)

    for col, (org, name, title_text) in enumerate([
        ("[CLIENT COMPANY PTE LTD]",
         "[Recipient Name]", "[Title]"),
        ("STARTECH INNOVATION PTE LTD",
         "Robert Rahardja", "Managing Director"),
    ]):
        cell = sig.cell(0, col)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(12)
        run(p, org, bold=True, italic=True, size=10, color=TEXT, font=FONT_SERIF)
        body(cell, "")
        body(cell, "")
        body(cell, "Signature: _______________________________")
        body(cell, name, bold=True)
        body(cell, title_text, color=MUTED, size=9)
        body(cell, "Date: ______________________", size=9)

    out = HERE / "Startech-Quotation-Template.docx"
    doc.save(out)
    return out


# ─── Invoice template ─────────────────────────────────────────────────────
def build_invoice():
    doc = Document()
    configure_base_styles(doc)
    setup_page(doc)
    build_header(doc)
    build_footer(doc)

    # Title
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    run(title, "INVOICE", size=26, color=TEXT, spacing=1.5)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run(subtitle, "[Project / engagement title]", size=11, color=MUTED,
        italic=True)

    # Meta + client block
    meta = doc.add_table(rows=1, cols=2)
    meta.autofit = False
    remove_all_table_borders(meta)
    meta.columns[0].width = Mm(60)
    meta.columns[1].width = Mm(106)

    left = meta.cell(0, 0)
    left.width = Mm(60)
    body(left, "INVOICE NO.", bold=True, size=9, color=MUTED)
    body(left, "STI-INV[YYMM]-[DDDD]-[NN]")
    body(left, "DATE", bold=True, size=9, color=MUTED)
    body(left, "[13 Jan 2026]")
    body(left, "REFERENCE", bold=True, size=9, color=MUTED)
    body(left, "[STI-Q2601-0108]")
    body(left, "")
    body(left, "+65 9069 3236")
    body(left, "robert@startech-innovation.com", color=BLUE)
    body(left, "7500A Beach Road, Singapore 199591", size=9, color=MUTED)

    right = meta.cell(0, 1)
    right.width = Mm(106)
    body(right, "BILL TO", bold=True, size=9, color=MUTED)
    body(right, "[CLIENT COMPANY PTE LTD]", bold=True)
    body(right, "[Street address, #floor-unit, Singapore 000000]")
    body(right, "")
    body(right, "ATTENTION", bold=True, size=9, color=MUTED)
    body(right, "[Dr. Recipient Name] — [Title]")

    h2(doc, "Project")
    body(doc, "[Project title, e.g. ReallyAgent Contextualization]")
    body(
        doc,
        "Per accepted quotation [STI-Q2601-0108] dated [8 Jan 2026].",
        italic=True, color=MUTED, size=9,
    )

    h2(doc, "Payment terms")
    body(doc, "Payment method: Bank transfer or as mutually agreed.")

    h2(doc, "Invoice details")
    styled_table(
        doc,
        headers=["Item", "Description", "Category", "Amount"],
        rows=[
            ["1",
             "[Line description — e.g. Initial Payment tranche 1]",
             "[Development Services]",
             "SGD 0,000"],
        ],
        col_widths_mm=[14, 90, 32, 30],
    )

    # Grand total
    total_tbl = doc.add_table(rows=1, cols=2)
    total_tbl.autofit = False
    remove_all_table_borders(total_tbl)
    total_tbl.columns[0].width = Mm(136)
    total_tbl.columns[1].width = Mm(30)
    left_t = total_tbl.cell(0, 0)
    right_t = total_tbl.cell(0, 1)
    left_t.width = Mm(136)
    right_t.width = Mm(30)
    set_cell_borders(
        left_t,
        top={"val": "single", "sz": 12, "color": "18181B"},
    )
    set_cell_borders(
        right_t,
        top={"val": "single", "sz": 12, "color": "18181B"},
    )
    p = left_t.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    run(p, "GRAND TOTAL", bold=True, size=11, color=TEXT,
        spacing=1.0, upper=True)
    p2 = right_t.paragraphs[0]
    p2.paragraph_format.space_before = Pt(3)
    run(p2, "SGD 0,000", bold=True, size=11, color=TEXT)

    h2(doc, "Notes")
    body(
        doc,
        "[Any milestone context, remaining tranches, or special "
        "instructions for this invoice.]",
    )

    h2(doc, "Banking details")
    banking = [
        ("Bank", "DBS Bank Limited"),
        ("Account name", "Robert Rahardja"),
        ("Account number", "001-0-059172"),
        ("SWIFT / BIC", "DBSSSGSG"),
        ("Bank address",
         "12 Marina Boulevard, DBS Asia Central, "
         "Marina Bay Financial Centre, Tower 3, Singapore 018982"),
    ]
    b_tbl = doc.add_table(rows=len(banking), cols=2)
    b_tbl.autofit = False
    remove_all_table_borders(b_tbl)
    b_tbl.columns[0].width = Mm(40)
    b_tbl.columns[1].width = Mm(126)
    for i, (k, v) in enumerate(banking):
        lc = b_tbl.cell(i, 0)
        rc = b_tbl.cell(i, 1)
        lc.width = Mm(40)
        rc.width = Mm(126)
        pk = lc.paragraphs[0]
        run(pk, k, bold=True, size=9, color=MUTED, upper=True, spacing=0.8)
        pv = rc.paragraphs[0]
        run(pv, v, size=10)

    h2(doc, "Terms & conditions")
    body(
        doc,
        "Payment to be made within 30 days via the payment method shown "
        "above. This invoice confirms your order.",
    )

    body(doc, "")
    body(doc, "Thank you for your business. "
              "It's a pleasure to work with you on your project.",
         italic=True)

    body(doc, "")
    run_p = doc.add_paragraph()
    run(run_p, "Robert Rahardja", bold=True, italic=True,
        font=FONT_SERIF, size=11)
    body(doc, "Managing Director, Startech Innovation Pte Ltd",
         color=MUTED, size=9)
    body(doc, "Direct: +65 9069 3236  |  "
              "Email: robert@startech-innovation.com",
         size=9, color=MUTED)

    out = HERE / "Startech-Invoice-Template.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    q = build_quotation()
    i = build_invoice()
    print(f"Wrote {q}")
    print(f"Wrote {i}")
